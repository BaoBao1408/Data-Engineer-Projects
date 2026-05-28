"""
Word Document Extractor – extract content from .docx files.
Handles paragraphs, headings, tables, embedded metadata.
Typical use: extract text from audit reports, compliance memos, SOPs.
"""
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from loguru import logger


@dataclass
class DocxTable:
    table_index: int
    rows: int
    cols: int
    data: list[list[str]]
    headers: list[str] = field(default_factory=list)

    def __post_init__(self):
        if self.data:
            self.headers = self.data[0]


@dataclass
class DocxDocument:
    file_path: str
    file_name: str
    file_size_bytes: int
    title: str
    author: str
    created: Optional[str]
    modified: Optional[str]
    sections: list[dict]       # {"heading": str, "level": int, "content": str}
    full_text: str
    tables: list[DocxTable]
    word_count: int = 0
    extraction_status: str = "success"
    error_message: Optional[str] = None

    def __post_init__(self):
        self.word_count = len(self.full_text.split())


class WordExtractor:
    """
    Extract structured content from Microsoft Word (.docx) documents.
    Preserves document structure (headings → sections → body text).
    """

    HEADING_STYLES = {
        "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
        "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
        "Title": 0,
    }

    def __init__(self, extract_tables: bool = True):
        self.extract_tables = extract_tables

    def extract(self, file_path: str | Path) -> DocxDocument:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")

        file_size = file_path.stat().st_size
        logger.info(f"Extracting Word doc: {file_path.name}")

        try:
            doc = Document(str(file_path))
            metadata = self._extract_metadata(doc)
            sections, full_text = self._extract_body(doc)
            tables = self._extract_tables(doc) if self.extract_tables else []

            return DocxDocument(
                file_path=str(file_path),
                file_name=file_path.name,
                file_size_bytes=file_size,
                title=metadata.get("title", file_path.stem),
                author=metadata.get("author", ""),
                created=metadata.get("created"),
                modified=metadata.get("modified"),
                sections=sections,
                full_text=full_text,
                tables=tables,
            )

        except Exception as exc:
            logger.error(f"Word extraction failed: {exc}")
            return DocxDocument(
                file_path=str(file_path),
                file_name=file_path.name,
                file_size_bytes=file_size,
                title=file_path.stem,
                author="",
                created=None,
                modified=None,
                sections=[],
                full_text="",
                tables=[],
                extraction_status="failed",
                error_message=str(exc),
            )

    def _extract_metadata(self, doc: Document) -> dict:
        props = doc.core_properties
        return {
            "title": props.title or "",
            "author": props.author or "",
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
            "subject": props.subject or "",
            "keywords": props.keywords or "",
        }

    def _extract_body(self, doc: Document) -> tuple[list[dict], str]:
        """
        Walk paragraphs and build a section tree.
        Returns (sections_list, full_text).
        """
        sections: list[dict] = []
        current_section: Optional[dict] = None
        text_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            heading_level = self.HEADING_STYLES.get(style_name)

            if heading_level is not None:
                # Save current section if exists
                if current_section:
                    sections.append(current_section)

                current_section = {
                    "heading": text,
                    "level": heading_level,
                    "content": "",
                    "paragraphs": [],
                }
                text_parts.append(f"\n{'#' * (heading_level + 1)} {text}\n")
            else:
                # Body paragraph
                if current_section:
                    current_section["content"] += text + "\n"
                    current_section["paragraphs"].append(text)
                text_parts.append(text)

        # Flush last section
        if current_section:
            sections.append(current_section)

        full_text = "\n".join(text_parts)
        return sections, full_text

    def _extract_tables(self, doc: Document) -> list[DocxTable]:
        tables = []
        for i, table in enumerate(doc.tables):
            try:
                rows_data = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_data.append(cells)

                if not rows_data:
                    continue

                # Deduplicate merged cells
                cleaned_rows = []
                for row in rows_data:
                    cleaned = []
                    prev = None
                    for cell in row:
                        if cell != prev:
                            cleaned.append(cell)
                        else:
                            cleaned.append("")
                        prev = cell
                    cleaned_rows.append(cleaned)

                tables.append(DocxTable(
                    table_index=i,
                    rows=len(cleaned_rows),
                    cols=len(cleaned_rows[0]) if cleaned_rows else 0,
                    data=cleaned_rows,
                ))
            except Exception as e:
                logger.warning(f"Table {i} extraction failed: {e}")

        return tables

    def extract_to_chunks(
        self, file_path: str | Path, chunk_size: int = 800, overlap: int = 150
    ) -> list[dict]:
        """Split document text into overlapping chunks for RAG."""
        doc = self.extract(file_path)
        if doc.extraction_status != "success":
            return []

        chunks = []
        # Chunk by section first, then by word count
        for section in doc.sections:
            section_text = f"{section['heading']}\n{section['content']}"
            words = section_text.split()

            for i in range(0, len(words), chunk_size - overlap):
                chunk_words = words[i: i + chunk_size]
                chunks.append({
                    "chunk_index": len(chunks),
                    "text": " ".join(chunk_words),
                    "word_count": len(chunk_words),
                    "source": doc.file_name,
                    "source_type": "docx",
                    "section_heading": section["heading"],
                    "metadata": {
                        "file_name": doc.file_name,
                        "title": doc.title,
                        "author": doc.author,
                    },
                })

        return chunks
